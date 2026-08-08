from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "test_report_docx_deps"))

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = ROOT / "docs" / "final_report" / "Smart_Study_Test_Report.docx"
LOG_DIR = ROOT / "docs" / "final_report" / "test_reports" / "logs"


def shade(cell, colour):
    tc_pr = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), colour)
    tc_pr.append(fill)


def set_cell_text(cell, value, bold=False, colour=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    if colour:
        from docx.shared import RGBColor
        r.font.color.rgb = RGBColor.from_string(colour)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, heading in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], heading, True, "FFFFFF")
        shade(table.rows[0].cells[index], "365F91")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if index == 0:
                shade(cells[index], "D9EAF7")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    return p


def paragraph(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_start and text.startswith(bold_start):
        p.add_run(bold_start).bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def collected_backend_tests():
    lines = (LOG_DIR / "backend_test_collection.txt").read_text(encoding="utf-8", errors="replace").splitlines()
    tests = [line.strip() for line in lines if line.strip().startswith("tests/") and "::" in line]
    counts = Counter(test.split("::", 1)[0].replace("tests/", "") for test in tests)
    return tests, counts


def build():
    tests, backend_counts = collected_backend_tests()
    doc = Document()
    section = doc.sections[0]
    section.page_height, section.page_width = Cm(29.7), Cm(21.0)
    section.top_margin = section.bottom_margin = Cm(2.2)
    section.left_margin = section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].font.name = "Times New Roman"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SMART STUDY\nTEST REPORT").bold = True
    title.runs[0].font.name = "Times New Roman"
    title.runs[0].font.size = Pt(22)
    sub = doc.add_paragraph("Flutter Mobile Application and FastAPI Backend")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for line in ("Project: Smart Study", "Report date: 5 August 2026", "Overall status: PASS WITH OPEN ITEMS"):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "PASS" in line:
            p.runs[0].bold = True
    doc.add_page_break()

    heading(doc, "1. Purpose and Scope")
    paragraph(doc, "This report records the automated checks completed for the Smart Study mobile application and its backend. The tests cover important user interface behaviour, data models, navigation, API contracts, study functions, exams, social features, notifications, security rules and deployment support. The results are based on the current local working copy of the project on 5 August 2026.")
    paragraph(doc, "The automated suites passed inside the tested scope. However, the report does not claim that every possible form of testing is finished. Two PostgreSQL integration tests were skipped because a separate TEST_DATABASE_URL was not configured, and the Flutter formatting check found files that need formatting. Manual device, usability, load and penetration testing were not evidenced during this run.")

    heading(doc, "2. Test Environment")
    add_table(doc, ["Item", "Recorded value"], [
        ("Flutter application", "Local working tree; reference commit 0a9f9d7"),
        ("Backend application", "Local working tree; reference commit 95adbd3"),
        ("Dart SDK", "3.12.2 (stable)"),
        ("Python", "3.14.6"),
        ("Frontend test framework", "Flutter test and Dart analyser"),
        ("Backend test framework", "pytest, Ruff and Python compileall"),
        ("Execution platform", "Windows development environment"),
    ])

    heading(doc, "3. Execution Summary")
    summary = [
        ("Flutter formatting", "FAIL", "116 files checked; 31 would be changed by the formatter."),
        ("Dart static analysis", "PASS", "No issues found in lib and test."),
        ("Flutter automated tests", "PASS", "44 tests passed."),
        ("Backend Ruff formatting", "PASS", "85 files already formatted."),
        ("Backend Ruff lint", "PASS", "All checks passed."),
        ("Backend compilation", "PASS", "app and tests compiled without errors."),
        ("Backend pytest", "PASS WITH SKIPS", "75 passed and 2 skipped; 77 tests collected."),
    ]
    add_table(doc, ["Check", "Status", "Actual result"], summary)
    paragraph(doc, "The first backend pytest attempt had two setup errors caused by permission restrictions in the Windows temporary folder. The same suite was run again with a backend-local temporary directory and completed with 75 passes and 2 expected skips. Therefore, those first errors are recorded as an environment problem rather than an application failure.")

    heading(doc, "4. Flutter Test Coverage")
    paragraph(doc, "The Flutter suite contains 15 test files and 44 individual tests. These tests mainly check widgets, compact layouts, routing behaviour and model parsing. They give repeatable evidence for the tested components, but they do not replace a complete end-to-end test on physical devices.")
    flutter_rows = [
        ("app_bottom_nav_compact_test.dart", "Compact bottom navigation behaviour"),
        ("chat_message_model_test.dart", "Chat message parsing and model behaviour"),
        ("dashboard_memory_plan_test.dart", "Dashboard memory and revision plan"),
        ("exam_contribution_screen_test.dart", "Friend exam contribution screen"),
        ("exam_dashboard_widget_test.dart", "Exam dashboard widget layout"),
        ("exam_model_test.dart", "Exam data model parsing"),
        ("exam_question_library_picker_test.dart", "Question picker and quota rules"),
        ("exam_shell_navigation_test.dart", "Exam route navigation"),
        ("home_performance_overview_test.dart", "Home performance overview"),
        ("notification_preferences_test.dart", "Notification preference parsing"),
        ("performance_dashboard_test.dart", "Performance dashboard model and UI"),
        ("push_notification_route_test.dart", "Push-notification route handling"),
        ("shell_navigation_routes_test.dart", "Shell navigation and back behaviour"),
        ("subject_card_compact_test.dart", "Compact subject card layout"),
        ("widget_test.dart", "Main bottom-navigation widget"),
    ]
    add_table(doc, ["Test file", "Main coverage"], flutter_rows)

    heading(doc, "5. Backend Test Coverage")
    paragraph(doc, f"pytest collected {len(tests)} backend tests from {len(backend_counts)} files. The suite checks service logic, API response contracts, validation, access rules, real-time events and supporting deployment functions.")
    backend_rows = [(name, count) for name, count in sorted(backend_counts.items())]
    add_table(doc, ["Backend test file", "Collected tests"], backend_rows)

    heading(doc, "6. Main Functional Results")
    functional = [
        ("Navigation and responsive UI", "PASS", "Compact bottom navigation, shell routes, exam routes and selected compact widgets behaved as expected."),
        ("Dashboard and performance", "PASS", "Score aggregation, streak rules, memory plans and Flutter dashboard models were checked."),
        ("Quiz and spaced revision", "PASS", "AI quiz input contracts, quality filters, question validation and revision interval rules passed."),
        ("Exam workflow", "PASS", "Exam lifecycle, contribution rules, security, question selection and solution release checks passed."),
        ("Friends, chat and real-time events", "PASS", "Friendship, messages, notifications, unread counts and Socket.IO-related service behaviour passed."),
        ("Push notifications", "PASS", "Payload shape, token rules and notification route handling passed in automated tests."),
        ("Security controls", "PASS IN TESTED SCOPE", "Ownership, visibility, sharing, rate limiting and selected content protection rules passed."),
        ("Real PostgreSQL integration", "NOT EXECUTED", "Two tests were skipped because TEST_DATABASE_URL was not configured."),
    ]
    add_table(doc, ["Area", "Result", "Explanation"], functional)

    heading(doc, "7. Defects, Warnings and Open Items")
    for text in [
        "Formatting: Dart formatting is not fully clean. Thirty-one of 116 checked files would be changed by dart format.",
        "Database integration: the registration, login, owned-subject CRUD and database-seed integration tests need a safe PostgreSQL test database before execution.",
        "Warnings: pytest completed with 281 warnings, mainly related to Windows asyncio behaviour and deprecations in dependencies and Firebase messaging APIs.",
        "Manual evidence: this run did not establish physical-device compatibility, actual FCM delivery, usability results, production load performance, penetration-test results or release testing on every supported platform.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    heading(doc, "8. Conclusion")
    paragraph(doc, "The automated results show that the main tested parts of Smart Study are working correctly. All 44 Flutter tests passed, while 75 backend tests passed and two database integration tests were skipped. Static analysis, backend linting, backend formatting and compilation also passed. This provides useful evidence that the implemented application is stable within the automated scope used for this report.")
    paragraph(doc, "Before describing the project as fully tested, the remaining formatting changes should be applied, the PostgreSQL integration tests should be run, and important manual checks should be recorded. A suitable final-report statement is: ‘The implemented system completed its planned automated unit, widget, contract and service tests successfully, with two environment-dependent database integration tests recorded as pending.’")

    heading(doc, "Appendix A: Executed Commands and Evidence Logs")
    commands = [
        ("dart format --output=none --set-exit-if-changed lib test", "flutter_format_check.txt"),
        ("dart analyze lib test", "dart_analyze.txt"),
        ("flutter test --reporter expanded", "flutter_test_expanded.txt"),
        ("python -m ruff format --check app tests", "backend_ruff_format.txt"),
        ("python -m ruff check app tests", "backend_ruff_lint.txt"),
        ("python -m compileall -q app tests", "backend_compileall.txt"),
        ("python -m pytest -q (successful retry)", "backend_pytest_retry2.txt"),
        ("python -m pytest --collect-only -q", "backend_test_collection.txt"),
    ]
    add_table(doc, ["Command", "Saved evidence"], commands)

    heading(doc, "Appendix B: Backend Test Inventory")
    for test in tests:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(re.sub(r"^tests/", "", test))
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    # Page numbers in the footer.
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Smart Study Test Report  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
