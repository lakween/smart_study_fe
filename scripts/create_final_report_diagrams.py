from __future__ import annotations

import math
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "report_docx_deps"))

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = ROOT / "docs" / "final_report"
IMAGE_DIR = OUT_DIR / "generated_diagrams"
DOCX_PATH = OUT_DIR / "Smart_Study_Diagrams_and_Explanations.docx"

W, H = 1800, 1050
BG = "#F7F9FC"
INK = "#172033"
BLUE = "#4F46E5"
PURPLE = "#7C3AED"
TEAL = "#0F9D8A"
ORANGE = "#E58A19"
RED = "#D1495B"
LINE = "#536079"
PALE_BLUE = "#E8EAFE"
PALE_TEAL = "#E4F6F2"
PALE_ORANGE = "#FFF1D9"
PALE_RED = "#FBE7EA"
WHITE = "#FFFFFF"


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_HEAD = font(29, True)
F_BODY = font(24)
F_SMALL = font(20)
F_SMALL_B = font(20, True)


def canvas(title: str):
    image = Image.new("RGB", (W, H), BG)
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((35, 25, W - 35, H - 25), radius=28, fill=WHITE, outline="#D9DFEA", width=3)
    draw.text((W // 2, 42), title, font=F_TITLE, fill=INK, anchor="mt")
    draw.line((90, 125, W - 90, 125), fill="#D9DFEA", width=3)
    return image, draw


def wrapped(draw, text, box, used_font=F_BODY, fill=INK, align="center", spacing=8):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    lines = []
    for paragraph in text.split("\n"):
        words, current = paragraph.split(), ""
        for word in words:
            trial = word if not current else current + " " + word
            if draw.textbbox((0, 0), trial, font=used_font)[2] <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        if not words:
            lines.append("")
    line_height = draw.textbbox((0, 0), "Ag", font=used_font)[3] + spacing
    total = len(lines) * line_height - spacing
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        if align == "left":
            x, anchor = x1 + 18, "lm"
        else:
            x, anchor = (x1 + x2) / 2, "mm"
        draw.text((x, y + line_height / 2), line, font=used_font, fill=fill, anchor=anchor)
        y += line_height


def box(draw, xy, text, fill=PALE_BLUE, outline=BLUE, radius=22, used_font=F_BODY):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=4)
    wrapped(draw, text, xy, used_font=used_font)


def arrow(draw, start, end, color=LINE, width=5, label=None):
    draw.line((*start, *end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.6, -2.6):
        p = (
            end[0] + length * math.cos(angle + delta),
            end[1] + length * math.sin(angle + delta),
        )
        draw.line((*end, *p), fill=color, width=width)
    if label:
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        bbox = draw.textbbox((mx, my), label, font=F_SMALL, anchor="mm")
        draw.rounded_rectangle((bbox[0] - 8, bbox[1] - 4, bbox[2] + 8, bbox[3] + 4), 8, fill=WHITE)
        draw.text((mx, my), label, font=F_SMALL, fill=INK, anchor="mm")


def save(image, name):
    path = IMAGE_DIR / name
    image.save(path, "PNG", dpi=(180, 180))
    return path


def use_case():
    im, d = canvas("Smart Study Use Case Overview")
    # Actor
    d.ellipse((70, 270, 150, 350), outline=INK, width=5)
    d.line((110, 350, 110, 525), fill=INK, width=5)
    d.line((110, 390, 40, 455), fill=INK, width=5)
    d.line((110, 390, 180, 455), fill=INK, width=5)
    d.line((110, 525, 45, 620), fill=INK, width=5)
    d.line((110, 525, 175, 620), fill=INK, width=5)
    d.text((110, 665), "Student", font=F_HEAD, fill=INK, anchor="mm")
    cases = [
        (300, 180, 700, 285, "Register, log in and manage profile"),
        (300, 330, 700, 435, "Manage subjects, topics and documents"),
        (300, 480, 700, 585, "Create, generate and practise quizzes"),
        (300, 630, 700, 735, "Create, contribute to and sit exams"),
        (1020, 180, 1450, 285, "Find friends and exchange messages"),
        (1020, 330, 1450, 435, "View dashboard and performance"),
        (1020, 480, 1450, 585, "Receive revision and exam reminders"),
        (1020, 630, 1450, 735, "Control visibility and copy content"),
    ]
    for x1, y1, x2, y2, text in cases:
        d.ellipse((x1, y1, x2, y2), fill=PALE_BLUE, outline=BLUE, width=4)
        wrapped(d, text, (x1, y1, x2, y2), used_font=F_SMALL_B)
    hub = (770, 370, 950, 560)
    d.ellipse(hub, fill=PURPLE, outline=PURPLE, width=4)
    wrapped(d, "Smart Study\nSystem", hub, used_font=F_HEAD, fill=WHITE)
    arrow(d, (185, 450), (300, 380))
    for x1, y1, x2, y2, _ in cases:
        sx = x2 if x2 < W / 2 else x1
        ex = hub[0] if x2 < W / 2 else hub[2]
        arrow(d, (sx, (y1 + y2) // 2), (ex, (hub[1] + hub[3]) // 2), width=3)
    d.text((W // 2, 930), "Main actions available to an authenticated learner", font=F_BODY, fill=LINE, anchor="mm")
    return save(im, "figure_3_1_use_case.png")


def architecture():
    im, d = canvas("High-Level System Architecture")
    box(d, (70, 330, 370, 610), "Flutter Mobile App\n\nMaterial 3 UI\nRiverpod state\nDio API client", PALE_BLUE, BLUE)
    box(d, (500, 240, 850, 700), "FastAPI Application\n\nREST routers\nDomain services\nRepositories\nSocket.IO server", "#EEE8FF", PURPLE)
    box(d, (1030, 160, 1350, 330), "PostgreSQL\nAuthoritative data", PALE_TEAL, TEAL)
    box(d, (1030, 420, 1350, 590), "Protected File Storage\nPDF and images", PALE_ORANGE, ORANGE)
    box(d, (1030, 680, 1350, 850), "AI Service\nQuiz generation", PALE_RED, RED)
    box(d, (1460, 220, 1720, 400), "Firebase Cloud\nMessaging", PALE_ORANGE, ORANGE, used_font=F_SMALL_B)
    box(d, (1460, 610, 1720, 790), "Scheduler\nWorker", PALE_TEAL, TEAL, used_font=F_SMALL_B)
    arrow(d, (370, 420), (500, 420), BLUE, label="HTTPS / JSON")
    arrow(d, (500, 570), (370, 570), PURPLE, label="Socket.IO")
    arrow(d, (850, 320), (1030, 245), TEAL, label="SQL")
    arrow(d, (850, 480), (1030, 505), ORANGE, label="files")
    arrow(d, (850, 630), (1030, 765), RED, label="prompt / result")
    arrow(d, (850, 365), (1460, 310), ORANGE, label="push request")
    arrow(d, (1460, 690), (1350, 280), TEAL, label="scheduled queries")
    d.text((W // 2, 930), "The mobile client uses one FastAPI backend for study, social, exam and notification functions.", font=F_BODY, fill=LINE, anchor="mm")
    return save(im, "figure_3_2_system_architecture.png")


def erd():
    im, d = canvas("Simplified Entity Relationship Diagram")
    entities = {
        "users": (60, 390, 290, 510),
        "subjects": (390, 170, 650, 290),
        "topics": (760, 170, 1020, 290),
        "documents": (1130, 170, 1420, 290),
        "quizzes": (390, 400, 650, 520),
        "questions": (760, 400, 1020, 520),
        "quiz attempts": (390, 660, 650, 780),
        "spaced repetitions": (760, 660, 1050, 780),
        "exams": (1130, 400, 1420, 520),
        "exam participants": (1490, 400, 1740, 520),
        "friendships": (60, 620, 290, 740),
        "direct messages": (60, 800, 290, 920),
        "notifications": (1130, 660, 1420, 780),
    }
    def link(a, b, label=None):
        aa, bb = entities[a], entities[b]
        p1 = ((aa[0] + aa[2]) // 2, (aa[1] + aa[3]) // 2)
        p2 = ((bb[0] + bb[2]) // 2, (bb[1] + bb[3]) // 2)
        arrow(d, p1, p2, LINE, width=3, label=label)
    link("users", "subjects")
    link("subjects", "topics")
    link("topics", "documents")
    link("subjects", "quizzes")
    link("quizzes", "questions")
    link("quizzes", "quiz attempts")
    link("quiz attempts", "spaced repetitions", "updates")
    link("topics", "spaced repetitions")
    link("questions", "exams", "selected")
    link("exams", "exam participants")
    link("users", "friendships")
    link("friendships", "direct messages")
    link("users", "notifications")
    for name, xy in entities.items():
        box(d, xy, name.replace(" ", "\n"), PALE_BLUE if name in {"subjects", "topics", "documents", "quizzes", "questions"} else PALE_TEAL, BLUE if name in {"subjects", "topics", "documents", "quizzes", "questions"} else TEAL, used_font=F_SMALL_B)
    d.text((W // 2, 970), "Lines represent key one-to-many relationships unless a different action is labelled.", font=F_SMALL, fill=LINE, anchor="mm")
    return save(im, "figure_3_3_simplified_erd.png")


def navigation():
    im, d = canvas("Mobile Application Navigation Flow")
    box(d, (60, 405, 290, 555), "Splash", PALE_BLUE, BLUE)
    box(d, (360, 280, 620, 410), "Login / Register", PALE_BLUE, BLUE)
    box(d, (360, 590, 620, 720), "Forgot Password", PALE_ORANGE, ORANGE)
    box(d, (720, 405, 990, 555), "Authenticated Shell", "#EEE8FF", PURPLE)
    tabs = [
        (1080, 155, 1350, 265, "Home Dashboard"),
        (1080, 305, 1350, 435, "Subjects\nTopics / Quizzes"),
        (1080, 485, 1350, 595, "Exams"),
        (1080, 645, 1350, 775, "Friends\nMessages"),
        (1450, 300, 1720, 430, "Profile\nSettings"),
    ]
    arrow(d, (290, 480), (360, 345), BLUE)
    arrow(d, (490, 410), (490, 590), ORANGE)
    arrow(d, (620, 345), (720, 470), PURPLE, label="valid session")
    for x1, y1, x2, y2, _ in tabs:
        arrow(d, (990, 480), (x1, (y1 + y2) // 2), LINE, width=3)
    for x1, y1, x2, y2, text in tabs:
        box(d, (x1, y1, x2, y2), text, PALE_TEAL, TEAL, used_font=F_SMALL_B)
    d.text((W // 2, 920), "The shell keeps the main navigation dock visible while focused creation and attempt screens open outside it.", font=F_SMALL, fill=LINE, anchor="mm")
    return save(im, "figure_3_4_navigation_flow.png")


def sequence_diagram(title, actors, messages, filename):
    im, d = canvas(title)
    margin, top, bottom = 150, 190, 900
    step = (W - 2 * margin) / (len(actors) - 1)
    xs = [int(margin + i * step) for i in range(len(actors))]
    for x, actor in zip(xs, actors):
        box(d, (x - 115, top - 45, x + 115, top + 45), actor, PALE_BLUE, BLUE, used_font=F_SMALL_B)
        d.line((x, top + 45, x, bottom), fill="#BBC3D1", width=3)
    y = 285
    gap = min(95, int((bottom - y) / max(1, len(messages))))
    for source, target, label, color in messages:
        x1, x2 = xs[source], xs[target]
        arrow(d, (x1, y), (x2, y), color, width=4, label=label)
        y += gap
    return save(im, filename)


def ai_flow():
    return sequence_diagram(
        "AI Quiz Generation Sequence",
        ["Student", "Flutter App", "FastAPI", "Text Extraction", "AI Service", "Database"],
        [
            (0, 1, "select content", BLUE),
            (1, 2, "send generation request", PURPLE),
            (2, 3, "extract source text", ORANGE),
            (3, 2, "clean text", ORANGE),
            (2, 4, "request MCQs", RED),
            (4, 2, "structured questions", RED),
            (2, 1, "return editable draft", PURPLE),
            (0, 1, "review and confirm", BLUE),
            (1, 2, "save accepted quiz", TEAL),
            (2, 5, "store quiz and questions", TEAL),
        ],
        "figure_4_1_ai_quiz_sequence.png",
    )


def spaced_flow():
    im, d = canvas("Quiz Scoring and Spaced-Repetition Flow")
    nodes = [
        (70, 390, 300, 540, "Start server\nquiz session", PALE_BLUE, BLUE),
        (370, 390, 600, 540, "Submit answers\nwith session ID", PALE_BLUE, BLUE),
        (670, 390, 900, 540, "Backend calculates\nscore", "#EEE8FF", PURPLE),
        (990, 390, 1220, 540, "Score at least\n60%?", PALE_ORANGE, ORANGE),
        (1310, 230, 1650, 370, "Advance interval\n1, 3, 7, 14 or 30 days", PALE_TEAL, TEAL),
        (1310, 590, 1650, 730, "Reset interval\nto 1 day", PALE_RED, RED),
    ]
    for x1, y1, x2, y2, text, fill, outline in nodes:
        box(d, (x1, y1, x2, y2), text, fill, outline, used_font=F_SMALL_B)
    for a, b in [((300, 465), (370, 465)), ((600, 465), (670, 465)), ((900, 465), (990, 465))]:
        arrow(d, a, b, LINE)
    arrow(d, (1220, 430), (1310, 300), TEAL, label="Yes")
    arrow(d, (1220, 505), (1310, 660), RED, label="No")
    box(d, (610, 760, 1190, 880), "Save next revision date and expose it to the dashboard, topic view and reminder scheduler", PALE_BLUE, BLUE, used_font=F_SMALL_B)
    arrow(d, (1480, 370), (1000, 760), TEAL)
    arrow(d, (1480, 730), (1080, 760), RED)
    return save(im, "figure_4_2_spaced_repetition_flow.png")


def exam_flow():
    return sequence_diagram(
        "Collaborative Friend Exam Workflow",
        ["Organizer", "FastAPI", "Friends", "Private Drafts", "Exam Attempt", "Results"],
        [
            (0, 1, "create friend exam", BLUE),
            (1, 2, "send invitations", PURPLE),
            (2, 1, "accept or decline", PURPLE),
            (0, 3, "submit own quota", TEAL),
            (2, 3, "submit equal quota", TEAL),
            (3, 1, "counts only; content private", ORANGE),
            (0, 1, "publish when all are ready", BLUE),
            (1, 4, "create stable shuffled attempts", RED),
            (2, 4, "autosave and submit answers", RED),
            (4, 5, "release after shared close", TEAL),
        ],
        "figure_4_3_collaborative_exam_sequence.png",
    )


def notification_flow():
    im, d = canvas("Notification and Real-Time Delivery Flow")
    box(d, (70, 380, 330, 560), "Domain Event or\nScheduler", PALE_ORANGE, ORANGE)
    box(d, (440, 380, 740, 560), "FastAPI Notification\nService", "#EEE8FF", PURPLE)
    box(d, (850, 190, 1160, 350), "PostgreSQL\nDurable history", PALE_TEAL, TEAL)
    box(d, (850, 440, 1160, 600), "Socket.IO\nForeground event", PALE_BLUE, BLUE)
    box(d, (850, 690, 1160, 850), "Firebase Cloud\nMessaging", PALE_ORANGE, ORANGE)
    box(d, (1320, 330, 1690, 700), "Flutter App\n\nNotification inbox\nForeground refresh\nBackground push\nDeep-link navigation", PALE_BLUE, BLUE)
    arrow(d, (330, 470), (440, 470), ORANGE, label="trigger")
    arrow(d, (740, 420), (850, 270), TEAL, label="commit first")
    arrow(d, (740, 490), (850, 520), BLUE, label="emit")
    arrow(d, (740, 545), (850, 770), ORANGE, label="best effort")
    arrow(d, (1160, 270), (1320, 410), TEAL, label="REST history")
    arrow(d, (1160, 520), (1320, 520), BLUE, label="live")
    arrow(d, (1160, 770), (1320, 620), ORANGE, label="push")
    d.text((W // 2, 930), "Database history remains available even if Socket.IO or Firebase delivery is temporarily unavailable.", font=F_SMALL, fill=LINE, anchor="mm")
    return save(im, "figure_4_4_notification_flow.png")


def deployment():
    im, d = canvas("Production Deployment Diagram")
    box(d, (70, 350, 350, 570), "Android / iOS\nSmart Study App", PALE_BLUE, BLUE)
    box(d, (470, 350, 750, 570), "Nginx\nTLS and reverse proxy", PALE_ORANGE, ORANGE)
    box(d, (870, 210, 1190, 390), "FastAPI systemd\nweb service", "#EEE8FF", PURPLE)
    box(d, (870, 550, 1190, 730), "Scheduler systemd\nworker", PALE_TEAL, TEAL)
    box(d, (1310, 180, 1690, 340), "PostgreSQL\ndatabase", PALE_TEAL, TEAL)
    box(d, (1310, 430, 1690, 590), "Shared protected\nupload storage", PALE_ORANGE, ORANGE)
    box(d, (1310, 680, 1690, 840), "Firebase Admin / FCM", PALE_RED, RED)
    arrow(d, (350, 460), (470, 460), BLUE, label="HTTPS")
    arrow(d, (750, 410), (870, 300), PURPLE, label="REST + Socket.IO")
    arrow(d, (1190, 270), (1310, 260), TEAL, label="SQL")
    arrow(d, (1190, 340), (1310, 510), ORANGE, label="files")
    arrow(d, (1190, 650), (1310, 760), RED, label="push")
    arrow(d, (1190, 640), (1310, 290), TEAL, label="scheduled queries")
    d.text((W // 2, 930), "REST and Socket.IO share the /smart-study origin; web and scheduler processes run as separate services.", font=F_SMALL, fill=LINE, anchor="mm")
    return save(im, "figure_4_5_deployment_diagram.png")


FIGURES = [
    (
        "3.1 Use Case Diagram",
        use_case,
        "Chapter 3 - Methodology and Artefact Design, after the functional requirements.",
        "Figure 3.1 gives a summary of the main actions supported by the Smart Study application. The student is the main actor of the system. After authentication, the student can organise learning materials, create or generate quizzes, practise, take part in exams, manage friendships, exchange text messages and view learning progress. Visibility and copying are also included because the application allows private, friends-only and public study content. This diagram keeps the scope clear without showing the lower-level API operations.",
        "Figure 3.1 shows the main functions that are available to a student in the Smart Study application.",
    ),
    (
        "3.2 System Architecture Diagram",
        architecture,
        "Chapter 3 - Methodology and Artefact Design, in the system architecture section.",
        "Figure 3.2 presents the main technical parts of Smart Study. The Flutter mobile application handles the screens and local user interaction. It communicates with the FastAPI backend through HTTPS requests, while Socket.IO is used for immediate foreground events. The backend applies the application rules and uses PostgreSQL as the main data store. Uploaded PDF and image files are kept in protected storage. AI quiz generation, background scheduling and Firebase push delivery are connected to the backend as supporting services. This design keeps important rules on the server instead of trusting the mobile device.",
        "As shown in Figure 3.2, the mobile application uses one FastAPI backend to access the database and supporting services.",
    ),
    (
        "3.3 Simplified Entity Relationship Diagram",
        erd,
        "Chapter 3 - Methodology and Artefact Design, in the database design section.",
        "Figure 3.3 shows the main data relationships used by the project. A user owns subjects, and each subject contains topics, documents and quizzes. Questions belong to quizzes and may also be included in exam papers. Quiz attempts store a learner's submitted work, and the result updates the spaced-repetition record. Exams are linked with participants, while friendships and direct messages connect users in the social part of the application. Notifications belong to individual users. The production database also includes supporting tables for sessions, refresh tokens, device tokens and collaborative contributions, but these were left out of the figure to keep it readable.",
        "The central relationships between the study, exam and social data are presented in Figure 3.3.",
    ),
    (
        "3.4 Application Navigation Flow",
        navigation,
        "Chapter 3 - Methodology and Artefact Design, after the user-interface design discussion.",
        "Figure 3.4 describes the high-level navigation of the Flutter application. A new or signed-out user moves from the splash screen to the authentication screens. A valid session opens the authenticated shell, which contains the Home, Subjects, Exams, Friends and Profile areas. The notification inbox and settings are reached from the relevant signed-in areas, although they are not separate bottom-navigation tabs. Detailed creation, contribution and active-attempt screens are opened as focused workflows, while normal detail and result pages keep the main navigation available where suitable.",
        "Figure 3.4 illustrates how authentication separates the public entry screens from the main application navigation.",
    ),
    (
        "4.1 AI Quiz Generation Sequence Diagram",
        ai_flow,
        "Chapter 4 - Development and Testing, in the AI quiz-generation implementation section.",
        "Figure 4.1 explains the AI quiz-generation process. The student first chooses learning content in the Flutter application. The request is sent to FastAPI, where source text is extracted and prepared before the AI service is contacted. The generated multiple-choice questions are returned as an editable draft instead of being saved immediately. This gives the student a chance to review and correct the content. Only the accepted quiz and questions are stored in the database. The review stage is important because AI output may contain unsuitable or unclear questions.",
        "The sequence in Figure 4.1 keeps the student in control by returning an editable draft before the quiz is saved.",
    ),
    (
        "4.2 Spaced-Repetition Flowchart",
        spaced_flow,
        "Chapter 4 - Development and Testing, in the quiz practice and memory scheduling section.",
        "Figure 4.2 shows how quiz results control the revision schedule. A practice attempt starts as a server-owned session, and the session identifier is sent back with the answers. The backend calculates the score instead of accepting a score from the mobile application. A score of 60 percent or more moves the learner to the next interval in the 1, 3, 7, 14 and 30-day sequence. A lower score resets the interval to one day. The calculated next revision date is stored and later shown in the dashboard, topic details and reminder process.",
        "As shown in Figure 4.2, the backend updates the revision interval using the result of the submitted quiz attempt.",
    ),
    (
        "4.3 Collaborative Exam Sequence Diagram",
        exam_flow,
        "Chapter 4 - Development and Testing, in the collaborative exam implementation section.",
        "Figure 4.3 presents the workflow of a friend exam. The organiser creates the draft and invites friends. Every accepted participant, including the organiser, contributes the same required number of questions. Contributions remain private, and the lobby displays progress counts without showing another person's questions. The organiser can publish only after the invitation and contribution conditions are satisfied. The backend then creates stable shuffled attempts, autosaves answers and releases the shared results after the exam has closed. This process provides equal participation while protecting the private question drafts.",
        "Figure 4.3 demonstrates the equal contribution and private question rules used by collaborative exams.",
    ),
    (
        "4.4 Notification Delivery Diagram",
        notification_flow,
        "Chapter 4 - Development and Testing, in the real-time and push-notification section.",
        "Figure 4.4 explains the three notification delivery paths. A feature event or scheduler task first creates a notification through the backend service. The record is committed to PostgreSQL, which remains the reliable notification history. Socket.IO provides an immediate update while the application is open, and Firebase Cloud Messaging provides best-effort delivery when the mobile app is in the background or closed. The Flutter application can still reload the notification inbox through the REST API if either live delivery method is unavailable.",
        "The layered delivery method in Figure 4.4 keeps PostgreSQL as the reliable history while Socket.IO and Firebase improve delivery speed.",
    ),
    (
        "4.5 Production Deployment Diagram",
        deployment,
        "Chapter 4 - Development and Testing, near the deployment or implementation-environment section.",
        "Figure 4.5 gives an overview of the production deployment. Mobile clients connect through HTTPS to Nginx, which terminates TLS and forwards REST and Socket.IO traffic to the FastAPI web service. The web service uses PostgreSQL and protected shared upload storage. A separate scheduler service performs exam lifecycle and revision reminder work without blocking normal API requests. Firebase Admin is used for push delivery when it is enabled. Separating the web and scheduler services makes their responsibilities clearer and allows them to be monitored independently.",
        "Figure 4.5 shows how the public mobile application reaches the separate FastAPI web and scheduler services.",
    ),
]


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def build_docx(paths):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    for style_name in ("Title", "Heading 1", "Heading 2"):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SMART STUDY\nREPORT DIAGRAMS AND EXPLANATIONS")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 41, 85)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared as supporting material for the Undergraduate Thesis Report")
    r.italic = True
    r.font.size = Pt(12)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Use of this document: ").bold = True
    note.add_run(
        "The diagrams can be copied into the recommended chapters of the final report. "
        "Figure numbers may be changed after the final report structure is confirmed. Each figure must be mentioned in the main text."
    )

    doc.add_heading("Diagram List", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Figure", "Diagram", "Recommended chapter"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        set_cell_shading(cell, "D9E2F3")
    for title, _, location, _, _ in FIGURES:
        row = table.add_row().cells
        number, name = title.split(" ", 1)
        row[0].text = number
        row[1].text = name
        row[2].text = location.split(",", 1)[0]

    doc.add_page_break()
    for index, ((title, _, location, explanation, reference_sentence), image_path) in enumerate(zip(FIGURES, paths)):
        doc.add_heading(title, level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(6.3))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = caption.add_run(f"Figure {title.split(' ', 1)[0]}: {title.split(' ', 1)[1]}")
        cr.italic = True
        cr.font.size = Pt(10)

        p = doc.add_paragraph()
        p.add_run("Recommended report location: ").bold = True
        p.add_run(location)
        doc.add_heading("Explanation", level=2)
        doc.add_paragraph(explanation)
        doc.add_heading("Example sentence for the report", level=2)
        p = doc.add_paragraph()
        r = p.add_run(reference_sentence)
        r.italic = True
        if index < len(paths) - 1:
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("Evidence and Editing Notes", level=1)
    notes = [
        "These diagrams were prepared from the current Flutter routes, Riverpod architecture, FastAPI services, PostgreSQL migrations and project documentation.",
        "The ER diagram is intentionally simplified. Use the migration files when a complete table-level diagram is required.",
        "Do not claim that every external service was performance-tested only because it appears in an architecture diagram.",
        "Update figure numbers, captions and chapter references after the final thesis table of contents is fixed.",
        "Keep each diagram close to the paragraph that explains it. A figure should not be included without discussion in the main text.",
    ]
    for item in notes:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(DOCX_PATH)


def main():
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    paths = [factory() for _, factory, _, _, _ in FIGURES]
    build_docx(paths)
    print(f"Created {len(paths)} diagrams in {IMAGE_DIR}")
    print(f"Created DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()
